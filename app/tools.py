# =============================================================================
# app/tools.py — Herramientas del agente AutoFlota-AI
# Todas las herramientas utilizadas por el agente están declaradas aquí.
# =============================================================================
from __future__ import annotations

import csv
import io
import json
import logging
import mimetypes
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import structlog
from google.adk.tools import ToolContext

from app.config import settings
from app.prompts import EXTRACTION_PROMPT, SUMMARY_PROMPT

logger = structlog.get_logger(__name__)

# =============================================================================
# SECCIÓN 1 — LECTURA DE DOCUMENTOS
# =============================================================================


def leer_archivo_pdf(ruta_archivo: str, tool_context: ToolContext | None = None) -> dict:
    """
    Lee y extrae el contenido de texto de un archivo PDF.

    Utiliza pdfplumber para documentos nativos y Google Document AI
    para documentos escaneados/imágenes.

    Args:
        ruta_archivo: Ruta local o URI de GCS (gs://bucket/path) al archivo PDF.
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con: texto_completo, paginas, metadatos, fuente.
    """
    try:
        import pdfplumber

        logger.info("Leyendo archivo PDF", ruta=ruta_archivo)

        # Si es URI de GCS, descargarlo primero
        ruta_local = _resolver_ruta_archivo(ruta_archivo)

        texto_paginas: list[str] = []
        metadatos: dict = {}

        with pdfplumber.open(ruta_local) as pdf:
            metadatos = {
                "total_paginas": len(pdf.pages),
                "autor": pdf.metadata.get("Author", ""),
                "creacion": str(pdf.metadata.get("CreationDate", "")),
                "titulo": pdf.metadata.get("Title", ""),
            }
            for num_pag, pagina in enumerate(pdf.pages, start=1):
                texto = pagina.extract_text() or ""
                tablas = pagina.extract_tables() or []
                tablas_str = _tablas_a_texto(tablas)
                texto_paginas.append(
                    f"--- Página {num_pag} ---\n{texto}\n{tablas_str}"
                )

        texto_completo = "\n\n".join(texto_paginas)

        # Si el texto es muy escaso, intentar con Document AI
        if len(texto_completo.strip()) < 200 and settings.document_ai_processor_id:
            logger.info("PDF escaneado detectado, usando Document AI")
            resultado_dai = procesar_con_document_ai(ruta_archivo, tool_context)
            texto_completo = resultado_dai.get("texto_completo", texto_completo)

        resultado = {
            "exito": True,
            "texto_completo": texto_completo,
            "paginas": texto_paginas,
            "metadatos": metadatos,
            "fuente": ruta_archivo,
            "tipo_archivo": "pdf",
        }

        if tool_context:
            tool_context.state["ultimo_documento_leido"] = resultado

        return resultado

    except Exception as exc:
        logger.error("Error al leer PDF", error=str(exc))
        return {"exito": False, "error": str(exc), "fuente": ruta_archivo}


def leer_archivo_excel(ruta_archivo: str, tool_context: ToolContext | None = None) -> dict:
    """
    Lee y extrae el contenido de un archivo Excel (.xlsx o .xls).

    Extrae todas las hojas, detecta tablas de flotilla y coberturas,
    y convierte los datos en texto estructurado para su análisis.

    Args:
        ruta_archivo: Ruta local o URI de GCS al archivo Excel.
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con texto_completo, hojas, metadatos, fuente.
    """
    try:
        import pandas as pd

        logger.info("Leyendo archivo Excel", ruta=ruta_archivo)
        ruta_local = _resolver_ruta_archivo(ruta_archivo)

        xl = pd.ExcelFile(ruta_local)
        hojas: dict[str, Any] = {}
        bloques_texto: list[str] = []

        for nombre_hoja in xl.sheet_names:
            df = pd.read_excel(ruta_local, sheet_name=nombre_hoja, header=None)
            df = df.fillna("")

            # Convertir a texto tabular legible
            texto_hoja = f"=== Hoja: {nombre_hoja} ===\n"
            texto_hoja += df.to_string(index=False, header=False)
            bloques_texto.append(texto_hoja)

            # Guardar datos crudos como lista de dicts
            df_named = pd.read_excel(ruta_local, sheet_name=nombre_hoja)
            df_named = df_named.fillna("").astype(str)
            hojas[nombre_hoja] = df_named.to_dict(orient="records")

        texto_completo = "\n\n".join(bloques_texto)

        resultado = {
            "exito": True,
            "texto_completo": texto_completo,
            "hojas": hojas,
            "nombres_hojas": xl.sheet_names,
            "metadatos": {"total_hojas": len(xl.sheet_names)},
            "fuente": ruta_archivo,
            "tipo_archivo": "excel",
        }

        if tool_context:
            tool_context.state["ultimo_documento_leido"] = resultado

        return resultado

    except Exception as exc:
        logger.error("Error al leer Excel", error=str(exc))
        return {"exito": False, "error": str(exc), "fuente": ruta_archivo}


def leer_archivo_word(ruta_archivo: str, tool_context: ToolContext | None = None) -> dict:
    """
    Lee y extrae el contenido de un archivo Word (.docx o .doc).

    Extrae párrafos, tablas y metadatos del documento Word.

    Args:
        ruta_archivo: Ruta local o URI de GCS al archivo Word.
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con texto_completo, parrafos, tablas, metadatos, fuente.
    """
    try:
        from docx import Document

        logger.info("Leyendo archivo Word", ruta=ruta_archivo)
        ruta_local = _resolver_ruta_archivo(ruta_archivo)

        doc = Document(ruta_local)
        parrafos: list[str] = []
        tablas_texto: list[str] = []

        # Extraer párrafos
        for parrafo in doc.paragraphs:
            texto = parrafo.text.strip()
            if texto:
                parrafos.append(texto)

        # Extraer tablas
        for idx, tabla in enumerate(doc.tables):
            filas = []
            for fila in tabla.rows:
                celdas = [celda.text.strip() for celda in fila.cells]
                filas.append(" | ".join(celdas))
            tablas_texto.append(f"--- Tabla {idx + 1} ---\n" + "\n".join(filas))

        texto_completo = "\n".join(parrafos)
        if tablas_texto:
            texto_completo += "\n\n" + "\n\n".join(tablas_texto)

        # Metadatos del documento
        props = doc.core_properties
        metadatos = {
            "autor": props.author or "",
            "titulo": props.title or "",
            "creacion": str(props.created or ""),
            "modificacion": str(props.modified or ""),
            "total_parrafos": len(parrafos),
            "total_tablas": len(doc.tables),
        }

        resultado = {
            "exito": True,
            "texto_completo": texto_completo,
            "parrafos": parrafos,
            "tablas": tablas_texto,
            "metadatos": metadatos,
            "fuente": ruta_archivo,
            "tipo_archivo": "word",
        }

        if tool_context:
            tool_context.state["ultimo_documento_leido"] = resultado

        return resultado

    except Exception as exc:
        logger.error("Error al leer Word", error=str(exc))
        return {"exito": False, "error": str(exc), "fuente": ruta_archivo}


# =============================================================================
# SECCIÓN 2 — GOOGLE DOCUMENT AI
# =============================================================================


def procesar_con_document_ai(
    ruta_archivo: str, tool_context: ToolContext | None = None
) -> dict:
    """
    Procesa un documento con Google Cloud Document AI para extracción
    avanzada de texto y entidades (especialmente útil para PDFs escaneados).

    Args:
        ruta_archivo: Ruta local o URI de GCS al documento.
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con texto_completo, entidades, confianza, fuente.
    """
    try:
        from google.cloud import documentai
        from google.api_core.client_options import ClientOptions

        logger.info("Procesando con Document AI", ruta=ruta_archivo)
        ruta_local = _resolver_ruta_archivo(ruta_archivo)

        client_options = ClientOptions(
            api_endpoint=f"{settings.document_ai_location}-documentai.googleapis.com"
        )
        client = documentai.DocumentProcessorServiceClient(client_options=client_options)

        with open(ruta_local, "rb") as f:
            contenido = f.read()

        mime_type = _detectar_mime_type(ruta_local)

        raw_document = documentai.RawDocument(content=contenido, mime_type=mime_type)
        nombre_procesador = client.processor_version_path(
            project=settings.google_cloud_project,
            location=settings.document_ai_location,
            processor=settings.document_ai_processor_id,
            processor_version=settings.document_ai_processor_version,
        )

        request = documentai.ProcessRequest(
            name=nombre_procesador, raw_document=raw_document
        )
        respuesta = client.process_document(request=request)
        documento = respuesta.document

        # Extraer entidades
        entidades: list[dict] = []
        for entidad in documento.entities:
            entidades.append(
                {
                    "tipo": entidad.type_,
                    "mencion": entidad.mention_text,
                    "confianza": round(entidad.confidence, 4),
                    "valor_normalizado": str(entidad.normalized_value),
                }
            )

        resultado = {
            "exito": True,
            "texto_completo": documento.text,
            "entidades": entidades,
            "total_paginas": len(documento.pages),
            "fuente": ruta_archivo,
            "procesador": settings.document_ai_processor_id,
        }

        if tool_context:
            tool_context.state["document_ai_resultado"] = resultado

        return resultado

    except Exception as exc:
        logger.warning("Document AI no disponible o falló, continuando sin él", error=str(exc))
        return {
            "exito": False,
            "texto_completo": "",
            "entidades": [],
            "error": str(exc),
            "fuente": ruta_archivo,
        }


# =============================================================================
# SECCIÓN 3 — EXTRACCIÓN ESTRUCTURADA CON GEMINI
# =============================================================================


def extraer_variables_solicitud(
    texto_documento: str, tool_context: ToolContext | None = None
) -> dict:
    """
    Extrae las variables estructuradas de una solicitud de seguro usando
    el modelo Gemini a través de Vertex AI.

    Analiza el texto del documento y produce un JSON estructurado con
    todos los campos relevantes de la solicitud.

    Args:
        texto_documento: Texto completo extraído del documento de solicitud.
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con las variables extraídas de la solicitud.
    """
    try:
        import vertexai
        from vertexai.generative_models import GenerativeModel, GenerationConfig

        logger.info("Extrayendo variables con Gemini", longitud_texto=len(texto_documento))

        vertexai.init(
            project=settings.vertexai_project,
            location=settings.vertexai_location,
        )

        modelo = GenerativeModel(
            model_name=settings.gemini_model_pro,
            generation_config=GenerationConfig(
                temperature=0.1,  # Baja temperatura para mayor precisión
                top_p=0.8,
                max_output_tokens=8192,
                response_mime_type="application/json",
            ),
        )

        prompt = EXTRACTION_PROMPT.format(document_text=texto_documento[:50000])
        respuesta = modelo.generate_content(prompt)

        # Parsear JSON
        texto_json = respuesta.text.strip()
        if texto_json.startswith("```json"):
            texto_json = texto_json[7:]
        if texto_json.startswith("```"):
            texto_json = texto_json[3:]
        if texto_json.endswith("```"):
            texto_json = texto_json[:-3]

        variables = json.loads(texto_json.strip())
        variables["timestamp_extraccion"] = datetime.now(timezone.utc).isoformat()
        variables["modelo_extraccion"] = settings.gemini_model_pro

        if tool_context:
            tool_context.state["variables_extraidas"] = variables

        logger.info("Variables extraídas exitosamente", total_vehiculos=
                    variables.get("flotilla", {}).get("total_vehiculos"))

        return {"exito": True, "variables": variables}

    except Exception as exc:
        logger.error("Error al extraer variables", error=str(exc))
        return {"exito": False, "error": str(exc), "variables": {}}


# =============================================================================
# SECCIÓN 4 — GENERACIÓN DE CSV
# =============================================================================


def generar_reporte_csv(
    variables: dict,
    ruta_salida: str | None = None,
    tool_context: ToolContext | None = None,
) -> dict:
    """
    Genera un archivo CSV estructurado con las variables principales
    de la solicitud de seguro de flotilla.

    El CSV incluye una fila por cada vehículo de la flotilla, con
    los datos generales de la solicitud repetidos en cada fila.

    Args:
        variables: Diccionario de variables extraídas por extraer_variables_solicitud.
        ruta_salida: Ruta donde guardar el CSV. Si es None, se genera en /tmp.
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con ruta_csv, contenido_csv, total_filas, exito.
    """
    try:
        logger.info("Generando reporte CSV")

        if not ruta_salida:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            cliente_nombre = variables.get("cliente", {}).get("nombre", "sin_nombre")
            cliente_nombre = "".join(c for c in cliente_nombre if c.isalnum() or c in "_-")[:30]
            ruta_salida = f"/tmp/solicitud_{cliente_nombre}_{timestamp}.csv"

        # Datos del encabezado (aplican a todos los vehículos)
        cliente = variables.get("cliente", {})
        broker = variables.get("broker", {})
        prima = variables.get("prima", {})
        fechas = variables.get("fechas", {})
        siniestros = variables.get("siniestros", {})
        flotilla = variables.get("flotilla", {})
        vehiculos = flotilla.get("vehiculos", [])

        # Coberturas como columnas planas
        coberturas = variables.get("coberturas", [])
        coberturas_dict = {}
        for cob in coberturas:
            nombre_cob = cob.get("nombre", "cobertura_desconocida")
            clave = f"cobertura_{nombre_cob.lower().replace(' ', '_')[:30]}_limite"
            coberturas_dict[clave] = cob.get("limite", "No especificado")
            clave_ded = f"cobertura_{nombre_cob.lower().replace(' ', '_')[:30]}_deducible"
            coberturas_dict[clave_ded] = cob.get("deducible", "No especificado")

        # Construir filas
        filas: list[dict] = []
        vehiculos_iter = vehiculos if vehiculos else [{}]

        for vehiculo in vehiculos_iter:
            fila = {
                # Identificación
                "nombre_cliente": cliente.get("nombre", "No especificado"),
                "rfc_cliente": cliente.get("rfc", "No especificado"),
                "giro_empresa": cliente.get("giro_empresa", "No especificado"),
                "nombre_broker": broker.get("nombre", "No especificado"),
                "agencia_broker": broker.get("agencia", "No especificado"),
                "clave_agente": broker.get("clave_agente", "No especificado"),
                # Coberturas
                **coberturas_dict,
                "otras_coberturas": ", ".join(
                    [c.get("nombre", "") for c in coberturas]
                ),
                # Prima
                "prima_maxima_esperada": prima.get("maxima_esperada", "No especificado"),
                "moneda": prima.get("moneda", "MXN"),
                "forma_pago": prima.get("forma_pago", "No especificado"),
                # Flotilla
                "total_vehiculos_flotilla": flotilla.get("total_vehiculos", "No especificado"),
                # Vehículo específico
                "numero_vehiculo": vehiculo.get("numero", ""),
                "marca": vehiculo.get("marca", "No especificado"),
                "modelo": vehiculo.get("modelo", "No especificado"),
                "año": vehiculo.get("año", "No especificado"),
                "version": vehiculo.get("version", "No especificado"),
                "placas": vehiculo.get("placas", "No especificado"),
                "numero_serie_vin": vehiculo.get("numero_serie", "No especificado"),
                "numero_motor": vehiculo.get("numero_motor", "No especificado"),
                "tipo_uso": vehiculo.get("tipo_uso", "No especificado"),
                "tipo_vehiculo": vehiculo.get("tipo_vehiculo", "No especificado"),
                "valor_comercial": vehiculo.get("valor_comercial", "No especificado"),
                "moneda_valor": vehiculo.get("moneda_valor", "MXN"),
                "conductores_habituales": vehiculo.get("conductores_habituales", "No especificado"),
                # Siniestros
                "periodo_siniestros_reportado": siniestros.get("periodo_reportado", "No especificado"),
                "numero_siniestros_3_anos": siniestros.get("total_siniestros", "No especificado"),
                "monto_total_siniestros": siniestros.get("monto_total_pagado", "No especificado"),
                "moneda_siniestros": siniestros.get("moneda_siniestros", "MXN"),
                # Fechas clave
                "fecha_devolucion_cotizacion": fechas.get("devolucion_cotizacion", "No especificado"),
                "fecha_inicio_vigencia": fechas.get("inicio_vigencia", "No especificado"),
                "fecha_fin_vigencia": fechas.get("fin_vigencia", "No especificado"),
                "fecha_vencimiento_cobertura_actual": fechas.get(
                    "vencimiento_cobertura_actual", "No especificado"
                ),
                "fecha_solicitud": fechas.get("fecha_solicitud", "No especificado"),
                # Condiciones y notas
                "condiciones_especiales": variables.get("condiciones_especiales", "No especificado"),
                "notas_broker": variables.get("notas_broker", "No especificado"),
                "alertas": "; ".join(variables.get("alertas", [])),
                "confianza_extraccion": variables.get("confianza_extraccion", "media"),
                "fuente_documento": variables.get("fuente_documento", "No especificado"),
                "timestamp_generacion": datetime.now(timezone.utc).isoformat(),
            }
            filas.append(fila)

        # Escribir CSV
        Path(ruta_salida).parent.mkdir(parents=True, exist_ok=True)
        with open(ruta_salida, "w", newline="", encoding="utf-8-sig") as f:
            if filas:
                escritor = csv.DictWriter(f, fieldnames=list(filas[0].keys()))
                escritor.writeheader()
                escritor.writerows(filas)

        # Generar también como string
        buffer = io.StringIO()
        if filas:
            escritor_buf = csv.DictWriter(buffer, fieldnames=list(filas[0].keys()))
            escritor_buf.writeheader()
            escritor_buf.writerows(filas)
        contenido_csv = buffer.getvalue()

        resultado = {
            "exito": True,
            "ruta_csv": ruta_salida,
            "contenido_csv": contenido_csv,
            "total_filas": len(filas),
            "total_columnas": len(filas[0]) if filas else 0,
            "nombre_archivo": Path(ruta_salida).name,
        }

        if tool_context:
            tool_context.state["csv_generado"] = resultado

        logger.info("CSV generado", ruta=ruta_salida, filas=len(filas))
        return resultado

    except Exception as exc:
        logger.error("Error al generar CSV", error=str(exc))
        return {"exito": False, "error": str(exc)}


# =============================================================================
# SECCIÓN 5 — BASE DE DATOS (FIRESTORE EN VERTEX AI)
# =============================================================================


def guardar_solicitud_base_de_datos(
    variables: dict,
    csv_ruta: str | None = None,
    metadata_adicional: dict | None = None,
    tool_context: ToolContext | None = None,
) -> dict:
    """
    Guarda el resultado del análisis de la solicitud en Firestore (Google Cloud).

    Persiste los datos estructurados de la solicitud como memoria a largo
    plazo para consultas futuras y análisis histórico.

    Args:
        variables: Variables extraídas de la solicitud.
        csv_ruta: Ruta al CSV generado (opcional).
        metadata_adicional: Metadatos adicionales a incluir (opcional).
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con id_documento, coleccion, exito.
    """
    try:
        from google.cloud import firestore

        logger.info("Guardando solicitud en Firestore")

        db = firestore.Client(
            project=settings.google_cloud_project,
            database=settings.firestore_database,
        )

        id_documento = str(uuid.uuid4())
        timestamp = datetime.now(timezone.utc)

        # Documento a guardar
        documento = {
            "id": id_documento,
            "timestamp_creacion": timestamp,
            "timestamp_actualizacion": timestamp,
            "nombre_cliente": variables.get("cliente", {}).get("nombre", ""),
            "rfc_cliente": variables.get("cliente", {}).get("rfc", ""),
            "nombre_broker": variables.get("broker", {}).get("nombre", ""),
            "total_vehiculos": variables.get("flotilla", {}).get("total_vehiculos"),
            "prima_maxima": variables.get("prima", {}).get("maxima_esperada"),
            "moneda": variables.get("prima", {}).get("moneda", "MXN"),
            "fecha_inicio_vigencia": variables.get("fechas", {}).get("inicio_vigencia"),
            "fecha_fin_vigencia": variables.get("fechas", {}).get("fin_vigencia"),
            "fecha_devolucion_cotizacion": variables.get("fechas", {}).get("devolucion_cotizacion"),
            "estatus": "NUEVA_SOLICITUD",
            "csv_ruta": csv_ruta or "",
            "datos_completos": variables,  # JSON completo
            **(metadata_adicional or {}),
        }

        # Guardar en colección de solicitudes
        col_ref = db.collection(settings.firestore_collection_solicitudes)
        doc_ref = col_ref.document(id_documento)
        doc_ref.set(documento)

        # También guardar/actualizar registro de cliente
        if variables.get("cliente", {}).get("nombre"):
            _actualizar_registro_cliente(db, variables.get("cliente", {}), id_documento)

        resultado = {
            "exito": True,
            "id_documento": id_documento,
            "coleccion": settings.firestore_collection_solicitudes,
            "proyecto": settings.google_cloud_project,
            "base_de_datos": settings.firestore_database,
        }

        if tool_context:
            tool_context.state["id_solicitud_guardada"] = id_documento
            tool_context.state["solicitud_guardada"] = True

        logger.info("Solicitud guardada en Firestore", id=id_documento)
        return resultado

    except Exception as exc:
        logger.error("Error al guardar en Firestore", error=str(exc))
        return {"exito": False, "error": str(exc)}


def buscar_historial_cliente(
    nombre_cliente: str | None = None,
    rfc_cliente: str | None = None,
    tool_context: ToolContext | None = None,
) -> dict:
    """
    Busca el historial de solicitudes y cotizaciones de un cliente en Firestore.

    Permite al agente contextualizar la solicitud actual con el historial
    previo del cliente para identificar patrones de siniestralidad o
    coberturas habituales.

    Args:
        nombre_cliente: Nombre del cliente a buscar.
        rfc_cliente: RFC del cliente (más preciso que el nombre).
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con total_solicitudes_previas, solicitudes, cliente_conocido.
    """
    try:
        from google.cloud import firestore

        logger.info("Buscando historial de cliente", nombre=nombre_cliente, rfc=rfc_cliente)

        db = firestore.Client(
            project=settings.google_cloud_project,
            database=settings.firestore_database,
        )

        col_ref = db.collection(settings.firestore_collection_solicitudes)
        solicitudes: list[dict] = []

        # Buscar por RFC (más preciso)
        if rfc_cliente:
            query = col_ref.where("rfc_cliente", "==", rfc_cliente).limit(10)
            docs = query.stream()
            for doc in docs:
                d = doc.to_dict()
                # Eliminar datos completos para resumir
                d.pop("datos_completos", None)
                solicitudes.append(d)

        # Si no encontró por RFC, buscar por nombre
        if not solicitudes and nombre_cliente:
            query = col_ref.where("nombre_cliente", "==", nombre_cliente).limit(10)
            docs = query.stream()
            for doc in docs:
                d = doc.to_dict()
                d.pop("datos_completos", None)
                solicitudes.append(d)

        resultado = {
            "exito": True,
            "cliente_conocido": len(solicitudes) > 0,
            "total_solicitudes_previas": len(solicitudes),
            "solicitudes": solicitudes,
        }

        if tool_context:
            tool_context.state["historial_cliente"] = resultado

        return resultado

    except Exception as exc:
        logger.warning("Error al buscar historial", error=str(exc))
        return {"exito": False, "cliente_conocido": False, "total_solicitudes_previas": 0,
                "solicitudes": [], "error": str(exc)}


# =============================================================================
# SECCIÓN 6 — GOOGLE CLOUD STORAGE
# =============================================================================


def subir_archivo_gcs(
    ruta_local: str,
    nombre_destino: str | None = None,
    bucket: str | None = None,
    carpeta: str = "outputs",
    tool_context: ToolContext | None = None,
) -> dict:
    """
    Sube un archivo local a Google Cloud Storage.

    Args:
        ruta_local: Ruta del archivo local a subir.
        nombre_destino: Nombre del blob en GCS. Si es None, usa el nombre del archivo.
        bucket: Nombre del bucket. Si es None, usa el bucket de outputs configurado.
        carpeta: Carpeta dentro del bucket (ej. 'outputs', 'submissions').
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con uri_gcs, url_publica, exito.
    """
    try:
        from google.cloud import storage

        bucket_name = bucket or settings.gcs_bucket_outputs
        if not nombre_destino:
            nombre_destino = Path(ruta_local).name

        blob_name = f"{carpeta}/{nombre_destino}" if carpeta else nombre_destino

        logger.info("Subiendo archivo a GCS", bucket=bucket_name, blob=blob_name)

        cliente = storage.Client(project=settings.google_cloud_project)
        bucket_obj = cliente.bucket(bucket_name)
        blob = bucket_obj.blob(blob_name)

        blob.upload_from_filename(ruta_local)
        uri_gcs = f"gs://{bucket_name}/{blob_name}"

        resultado = {
            "exito": True,
            "uri_gcs": uri_gcs,
            "bucket": bucket_name,
            "blob_name": blob_name,
            "tamaño_bytes": Path(ruta_local).stat().st_size,
        }

        if tool_context:
            tool_context.state["ultimo_archivo_gcs"] = uri_gcs

        logger.info("Archivo subido a GCS", uri=uri_gcs)
        return resultado

    except Exception as exc:
        logger.error("Error al subir a GCS", error=str(exc))
        return {"exito": False, "error": str(exc)}


def obtener_archivo_gcs(
    uri_gcs: str,
    directorio_local: str = "/tmp",
    tool_context: ToolContext | None = None,
) -> dict:
    """
    Descarga un archivo desde Google Cloud Storage a una ruta local.

    Args:
        uri_gcs: URI del archivo en GCS (ej. gs://bucket/path/file.pdf).
        directorio_local: Directorio local donde guardar el archivo.
        tool_context: Contexto del agente ADK (inyectado automáticamente).

    Returns:
        Diccionario con ruta_local, exito.
    """
    try:
        from google.cloud import storage

        if not uri_gcs.startswith("gs://"):
            return {"exito": False, "error": "URI debe comenzar con gs://"}

        partes = uri_gcs[5:].split("/", 1)
        bucket_name, blob_name = partes[0], partes[1]
        nombre_archivo = Path(blob_name).name
        ruta_local = str(Path(directorio_local) / nombre_archivo)

        logger.info("Descargando archivo de GCS", uri=uri_gcs, destino=ruta_local)

        Path(directorio_local).mkdir(parents=True, exist_ok=True)
        cliente = storage.Client(project=settings.google_cloud_project)
        bucket_obj = cliente.bucket(bucket_name)
        blob = bucket_obj.blob(blob_name)
        blob.download_to_filename(ruta_local)

        resultado = {"exito": True, "ruta_local": ruta_local, "uri_gcs": uri_gcs}

        if tool_context:
            tool_context.state["ultimo_archivo_local"] = ruta_local

        return resultado

    except Exception as exc:
        logger.error("Error al descargar de GCS", error=str(exc))
        return {"exito": False, "error": str(exc)}


# =============================================================================
# SECCIÓN 7 — UTILIDADES PRIVADAS
# =============================================================================


def _resolver_ruta_archivo(ruta: str) -> str:
    """Descarga el archivo si es una URI de GCS; devuelve la ruta local."""
    if ruta.startswith("gs://"):
        resultado = obtener_archivo_gcs(ruta)
        if resultado.get("exito"):
            return resultado["ruta_local"]
        raise FileNotFoundError(f"No se pudo descargar el archivo: {ruta}")
    if not Path(ruta).exists():
        raise FileNotFoundError(f"Archivo no encontrado: {ruta}")
    return ruta


def _tablas_a_texto(tablas: list) -> str:
    """Convierte tablas de pdfplumber a texto legible."""
    if not tablas:
        return ""
    bloques = []
    for idx, tabla in enumerate(tablas):
        filas = [" | ".join(str(celda or "") for celda in fila) for fila in tabla]
        bloques.append(f"[Tabla {idx + 1}]\n" + "\n".join(filas))
    return "\n\n".join(bloques)


def _detectar_mime_type(ruta: str) -> str:
    """Detecta el MIME type de un archivo por extensión."""
    mime, _ = mimetypes.guess_type(ruta)
    if mime:
        return mime
    ext = Path(ruta).suffix.lower()
    mapa = {
        ".pdf": "application/pdf",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
    }
    return mapa.get(ext, "application/octet-stream")


def _actualizar_registro_cliente(db: Any, cliente: dict, id_solicitud: str) -> None:
    """Crea o actualiza el registro de cliente en Firestore."""
    try:
        clave = cliente.get("rfc") or cliente.get("nombre", "").replace(" ", "_").lower()
        if not clave:
            return
        ref = db.collection(settings.firestore_collection_clientes).document(clave)
        doc = ref.get()
        if doc.exists:
            datos = doc.to_dict()
            solicitudes_prev = datos.get("solicitudes", [])
            solicitudes_prev.append(id_solicitud)
            ref.update({
                "solicitudes": solicitudes_prev,
                "timestamp_ultima_solicitud": datetime.now(timezone.utc),
                "total_solicitudes": len(solicitudes_prev),
            })
        else:
            ref.set({
                **cliente,
                "solicitudes": [id_solicitud],
                "timestamp_primer_solicitud": datetime.now(timezone.utc),
                "timestamp_ultima_solicitud": datetime.now(timezone.utc),
                "total_solicitudes": 1,
            })
    except Exception as exc:
        logger.warning("No se pudo actualizar registro de cliente", error=str(exc))


# =============================================================================
# SECCIÓN 8 — LISTA EXPORTABLE DE HERRAMIENTAS
# =============================================================================

# Lista de todas las herramientas disponibles para el agente
TODAS_LAS_HERRAMIENTAS = [
    leer_archivo_pdf,
    leer_archivo_excel,
    leer_archivo_word,
    procesar_con_document_ai,
    extraer_variables_solicitud,
    generar_reporte_csv,
    guardar_solicitud_base_de_datos,
    buscar_historial_cliente,
    subir_archivo_gcs,
    obtener_archivo_gcs,
]
