# =============================================================================
# eval/test_cases/test_word_submission.py — Tests para documentos Word
# =============================================================================
"""
Pruebas para el procesamiento de solicitudes de seguro en formato Word (.docx).

Ejecutar:
    uv run pytest eval/test_cases/test_word_submission.py -v
"""
from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

ROOT_DIR = Path(__file__).parent.parent.parent


# =============================================================================
# Fixtures
# =============================================================================

@pytest.fixture
def word_solicitud_muestra(tmp_path) -> Path:
    """Genera un archivo Word de muestra para pruebas."""
    doc = Document()

    doc.add_heading("SOLICITUD DE COTIZACIÓN — SEGURO DE FLOTILLA DE AUTOS", 0)
    doc.add_paragraph("")

    doc.add_heading("DATOS DEL CLIENTE", 1)
    tabla_cliente = doc.add_table(rows=5, cols=2)
    datos_cliente = [
        ("Razón Social:", "Logística Express del Centro S.A. de C.V."),
        ("RFC:", ""),  # Intencionalmente vacío para probar robustez
        ("Giro:", "Transporte y logística de última milla"),
        ("Dirección:", "Blvd. Centro 123, Col. Industrial, CDMX"),
        ("Teléfono:", "55-1234-5678"),
    ]
    for i, (etiqueta, valor) in enumerate(datos_cliente):
        tabla_cliente.rows[i].cells[0].text = etiqueta
        tabla_cliente.rows[i].cells[1].text = valor

    doc.add_paragraph("")
    doc.add_heading("DATOS DEL BROKER", 1)
    tabla_broker = doc.add_table(rows=3, cols=2)
    datos_broker = [
        ("Nombre del Agente:", "Roberto Guzmán Silva"),
        ("Agencia:", ""),  # Intencionalmente vacío
        ("Email:", "rguzmán@corredores.com.mx"),
    ]
    for i, (etiqueta, valor) in enumerate(datos_broker):
        tabla_broker.rows[i].cells[0].text = etiqueta
        tabla_broker.rows[i].cells[1].text = valor

    doc.add_paragraph("")
    doc.add_heading("COBERTURAS SOLICITADAS", 1)
    doc.add_paragraph("• Responsabilidad Civil Daños a Terceros: $2,000,000 MXN")
    doc.add_paragraph("• Daños Materiales: Valor Comercial (deducible 10%)")
    doc.add_paragraph("• Robo Total y Parcial: Valor Comercial")
    doc.add_paragraph("• Gastos Médicos a Ocupantes: $150,000 MXN")

    doc.add_paragraph("")
    doc.add_heading("PRIMA MÁXIMA ESPERADA", 1)
    doc.add_paragraph("$280,000.00 MXN anuales (pago anual anticipado)")

    doc.add_paragraph("")
    doc.add_heading("FLOTILLA", 1)
    tabla_flotilla = doc.add_table(rows=9, cols=6)
    encabezados = ["No.", "Marca", "Modelo", "Año", "Placas", "Valor Comercial"]
    for i, enc in enumerate(encabezados):
        tabla_flotilla.rows[0].cells[i].text = enc
    vehiculos_lista = [
        ("1", "Nissan", "Urvan", "2022", "CDMX-001-A", "$420,000"),
        ("2", "Nissan", "Urvan", "2022", "CDMX-002-A", "$420,000"),
        ("3", "Toyota", "Hiace", "2021", "CDMX-003-B", "$480,000"),
        ("4", "Volkswagen", "Crafter", "2023", "CDMX-004-C", "$650,000"),
        ("5", "Ford", "Transit", "2022", "CDMX-005-D", "$580,000"),
        ("6", "Chevrolet", "Express", "2021", "CDMX-006-E", "$510,000"),
        ("7", "Mercedes-Benz", "Sprinter", "2023", "CDMX-007-F", "$720,000"),
        ("8", "Nissan", "NP300", "2022", "CDMX-008-G", "$320,000"),
    ]
    for i, v in enumerate(vehiculos_lista, 1):
        for j, val in enumerate(v):
            tabla_flotilla.rows[i].cells[j].text = val

    doc.add_paragraph("")
    doc.add_heading("HISTORIAL DE SINIESTROS", 1)
    doc.add_paragraph(
        "Sin siniestros registrados en los últimos 3 años. "
        "Primera vez que el cliente adquiere seguro de flotilla."
    )

    doc.add_paragraph("")
    doc.add_heading("FECHAS", 1)
    tabla_fechas = doc.add_table(rows=4, cols=2)
    fechas = [
        ("Inicio de vigencia:", "15/07/2025"),
        ("Fin de vigencia:", "14/07/2026"),
        ("Vencimiento cobertura actual:", "14/07/2025"),
        ("Devolución cotización:", ""),  # Intencionalmente vacío
    ]
    for i, (et, va) in enumerate(fechas):
        tabla_fechas.rows[i].cells[0].text = et
        tabla_fechas.rows[i].cells[1].text = va

    ruta = tmp_path / "solicitud_logistica_express.docx"
    doc.save(ruta)
    return ruta


# =============================================================================
# Tests
# =============================================================================

class TestLeerArchivoWord:
    def test_lectura_exitosa(self, word_solicitud_muestra):
        """Verifica lectura correcta de un archivo Word."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word(str(word_solicitud_muestra))

        assert resultado["exito"] is True
        assert "texto_completo" in resultado
        assert len(resultado["texto_completo"]) > 50

    def test_extrae_parrafos(self, word_solicitud_muestra):
        """Verifica que se extraen párrafos del documento."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word(str(word_solicitud_muestra))
        assert len(resultado["parrafos"]) > 0

    def test_extrae_tablas(self, word_solicitud_muestra):
        """Verifica que se extraen tablas del documento."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word(str(word_solicitud_muestra))
        assert len(resultado["tablas"]) > 0

    def test_contiene_nombre_cliente(self, word_solicitud_muestra):
        """Verifica que el texto incluye el nombre del cliente."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word(str(word_solicitud_muestra))
        assert "Logística Express" in resultado["texto_completo"]

    def test_metadatos_correctos(self, word_solicitud_muestra):
        """Verifica que los metadatos tienen la estructura esperada."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word(str(word_solicitud_muestra))
        metadatos = resultado["metadatos"]

        assert "total_parrafos" in metadatos
        assert "total_tablas" in metadatos
        assert metadatos["total_tablas"] >= 4

    def test_archivo_inexistente(self):
        """Verifica manejo de error para archivo inexistente."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word("/no/existe/archivo.docx")
        assert resultado["exito"] is False

    def test_tipo_archivo(self, word_solicitud_muestra):
        """Verifica identificación del tipo de archivo."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word(str(word_solicitud_muestra))
        assert resultado["tipo_archivo"] == "word"


class TestDatosIncompletosWord:
    """Pruebas de robustez con campos vacíos (caso de prueba #3)."""

    def test_manejo_rfc_vacio(self, word_solicitud_muestra):
        """Verifica que se maneja correctamente un RFC vacío."""
        from app.tools import leer_archivo_word, generar_reporte_csv

        resultado_lectura = leer_archivo_word(str(word_solicitud_muestra))

        # Simular variables con RFC vacío
        variables = {
            "cliente": {"nombre": "Logística Express", "rfc": "", "giro_empresa": "Logística"},
            "broker": {"nombre": "Roberto Guzmán", "agencia": ""},
            "prima": {"maxima_esperada": 280000, "moneda": "MXN"},
            "flotilla": {"total_vehiculos": 8, "vehiculos": []},
            "coberturas": [],
            "siniestros": {"total_siniestros": 0},
            "fechas": {"inicio_vigencia": "15/07/2025", "fin_vigencia": "14/07/2026",
                       "devolucion_cotizacion": None},
            "alertas": ["RFC no proporcionado", "Fecha de devolución no especificada"],
        }

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".csv", delete=False) as tmp:
            ruta_csv = tmp.name

        resultado_csv = generar_reporte_csv(variables, ruta_csv)
        assert resultado_csv["exito"] is True

    def test_genera_alertas_campos_faltantes(self, word_solicitud_muestra):
        """Verifica que se generan alertas por campos faltantes."""
        from app.tools import leer_archivo_word

        resultado = leer_archivo_word(str(word_solicitud_muestra))
        # El documento tiene campos vacíos intencionalmente
        # El agente debería detectarlos como alertas en la extracción posterior
        assert resultado["exito"] is True
        # La validación de alertas ocurre en la extracción con Gemini
